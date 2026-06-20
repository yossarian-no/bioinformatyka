#!/usr/bin/env python3

import argparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_OUTPUT = "sprawozdanie.docx"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)

    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_margins(cell)
            if widths:
                cell.width = Cm(widths[index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

    document.add_paragraph()
    return table


def add_bullet(document, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    return document.add_paragraph(text, style=style)


def add_number(document, text):
    return document.add_paragraph(text, style="List Number")


def add_manual_number(document, number, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.first_line_indent = Cm(-0.65)
    paragraph.add_run(f"{number}. ").bold = True
    paragraph.add_run(text)
    return paragraph


def add_code_line(document, text):
    paragraph = document.add_paragraph(style="Kod")
    paragraph.add_run(text)
    return paragraph


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13)):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(26)
    styles["Title"].font.bold = True

    if "Kod" not in styles:
        code = styles.add_style("Kod", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Kod"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.8)
    code.paragraph_format.space_after = Pt(3)


def build_document(output_path=DEFAULT_OUTPUT):
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Sprawozdanie")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(
        "Sekwencjonowanie przez hybrydyzację\n"
        "Binarny chip z błędami negatywnymi\n"
        "(wariant z błędami negatywnymi)"
    )
    run.bold = True
    run.font.size = Pt(16)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(
        "Bioinformatyka, laboratoria L7\n\n"
        "Vasil Kusmartsev 156202\n"
        "Mateusz Kaźmierczak 160162"
    )
    document.add_page_break()

    document.add_heading("1. Opis i formalizacja problemu", level=1)
    document.add_paragraph(
        "Sekwencjonowanie przez hybrydyzację (SBH) polega na odtworzeniu "
        "sekwencji DNA długości n na podstawie spektrum krótkich sond "
        "(oligonukleotydów), z którymi hybrydyzowała badana sekwencja."
    )

    document.add_heading("1.1 Binarny chip", level=2)
    document.add_paragraph(
        "Binarny chip wykorzystuje również symbole niejednoznaczne IUPAC. "
        "Symbol jest zgodny z nukleotydem, jeżeli zbiór reprezentowanych "
        "przez niego zasad zawiera ten nukleotyd."
    )
    add_table(
        document,
        ["Symbol", "Zgodne nukleotydy"],
        [
            ("R", "A, G"),
            ("Y", "C, T"),
            ("S", "G, C"),
            ("W", "A, T"),
            ("N / Z / P", "A, C, G, T"),
        ],
        [4, 10],
    )
    document.add_paragraph(
        "Zgodność dwóch symboli zachodzi wtedy, gdy przecięcie ich zbiorów "
        "dopuszczalnych nukleotydów jest niepuste."
    )

    document.add_heading("1.2 Błędy negatywne (Err−)", level=2)
    document.add_paragraph(
        "Błąd negatywny oznacza brak części prawidłowych sond w spektrum. "
        "W konsekwencji rozwiązanie nie musi tworzyć ciągłego przejścia przez "
        "wszystkie oczekiwane sondy, a kryterium jakości stanowi liczba "
        "różnych sond pokrytych przez otrzymaną sekwencję."
    )

    document.add_heading("1.3 Formalizacja", level=2)
    document.add_paragraph("Dane wejściowe:")
    add_bullet(document, "n — długość szukanej sekwencji DNA,")
    add_bullet(document, "k — długość sondy,")
    add_bullet(document, "s₀ — znany fragment początkowy sekwencji,")
    add_bullet(document, "S = {s₁, …, sₘ} — zbiór unikalnych sond.")
    document.add_paragraph(
        "Szukana jest sekwencja d ∈ {A, C, G, T}ⁿ rozpoczynająca się od s₀, "
        "która maksymalizuje liczbę sond z S występujących w d zgodnie z "
        "regułami kompatybilności IUPAC. Problem jest NP-trudny."
    )

    document.add_heading("2. Opis podejścia dokładnego", level=1)
    document.add_paragraph(
        "Algorytm dokładny z pliku exact_b_m.py jest bezpośrednim modelem "
        "programowania całkowitoliczbowego (ILP). Model nie opiera się już "
        "wyłącznie na zgodności kolejnych sond parami, ponieważ taka zgodność "
        "nie gwarantuje istnienia jednej wspólnej konkretyzacji dla kilku "
        "nakładających się symboli IUPAC."
    )

    document.add_heading("2.1 Zmienne decyzyjne", level=2)
    add_bullet(
        document,
        "bₚ,ₐ = 1, gdy na pozycji p sekwencji wybrano nukleotyd "
        "a ∈ {A, C, G, T}.",
    )
    add_bullet(
        document,
        "zᵢ,ₜ = 1, gdy sonda i rozpoczyna się w wyniku na pozycji t.",
    )
    add_bullet(
        document,
        "yᵢ = 1, gdy sonda i jest pokryta przez wynikową sekwencję.",
    )

    document.add_heading("2.2 Funkcja celu i ograniczenia", level=2)
    document.add_paragraph("Funkcja celu maksymalizuje liczbę pokrytych sond:")
    add_code_line(document, "max  Σᵢ yᵢ")
    add_number(
        document,
        "Na każdej pozycji sekwencji wybrany jest dokładnie jeden nukleotyd.",
    )
    add_number(
        document,
        "Pozycje należące do znanego początku s₀ są ustalone.",
    )
    add_number(
        document,
        "Jeżeli zᵢ,ₜ = 1, każda pozycja sondy i musi być zgodna z "
        "nukleotydem wyniku na odpowiedniej pozycji.",
    )
    add_number(
        document,
        "Dla każdej sondy yᵢ jest równe sumie jej wybranych położeń; "
        "jedna sonda jest liczona najwyżej raz.",
    )
    document.add_paragraph(
        "Jeżeli program rozwiązujący CBC kończy pracę ze statusem „Optimal” "
        "(optymalny), rozwiązanie ma udowodnioną optymalność. Wynik zwrócony "
        "po limicie czasu może być poprawnym rozwiązaniem dopuszczalnym, ale "
        "nie wolno wtedy deklarować gwarancji optimum."
    )

    document.add_heading("2.3 Złożoność i zakres stosowania", level=2)
    add_table(
        document,
        ["Element modelu", "Rząd wielkości"],
        [
            ("Zmienne nukleotydów", "4n"),
            ("Potencjalne położenia sond", "O(m(n − k + 1))"),
            ("Ograniczenia zgodności", "O(mk(n − k + 1))"),
            ("Klasa problemu", "NP-trudny"),
        ],
        [7, 7],
    )
    document.add_paragraph(
        "Algorytm dokładny jest przeznaczony do małych instancji oraz do "
        "walidacji jakości heurystyki. Dla dużych danych liczba potencjalnych "
        "położeń sond szybko rośnie."
    )

    document.add_heading("3. Opis podejścia heurystycznego", level=1)
    document.add_paragraph(
        "Algorytm heurystyczny z pliku heur_b_m.py buduje rozwiązanie "
        "zachłannie, wykonuje losowe ponowne uruchomienia i stosuje lokalne mutacje "
        "punktowe. Znany fragment początkowy jest zachowywany również podczas "
        "przeszukiwania lokalnego."
    )

    document.add_heading("3.1 Konstrukcja zachłanna", level=2)
    document.add_paragraph(
        "W każdym kroku wybierana jest niepokryta sonda o najmniejszym koszcie "
        "dołączenia, czyli wymagająca dodania najmniejszej liczby nowych "
        "nukleotydów. Remisy rozstrzyga ocena liczby możliwych dobrych "
        "następników."
    )

    document.add_heading("3.2 Ponowne uruchomienia i poprawa lokalna", level=2)
    document.add_paragraph(
        "Po przebiegu deterministycznym wykonywane są przebiegi losowe. "
        "Następnie algorytm próbuje mutacji pojedynczych pozycji i zachowuje "
        "zmiany, które nie zmniejszają liczby pokrytych sond. Całość podlega "
        "limitowi czasu 95 s."
    )

    document.add_heading("3.3 Gwarancje", level=2)
    document.add_paragraph(
        "Heurystyka zawsze zwraca sekwencję wymaganej długości, ale nie daje "
        "dowodu optymalności. Jej wynik należy oceniać przez liczbę pokrytych "
        "sond i porównywać z optimum tylko na instancjach, dla których optimum "
        "zostało niezależnie wyznaczone."
    )

    document.add_heading("4. Analiza uzyskanych wyników", level=1)
    document.add_heading("4.1 Dane i procedura testowa", level=2)
    document.add_paragraph(
        "Plik input.xml zawiera instancję n = 500, k = 19, znany początek "
        "długości 19 oraz 864 unikalne sondy. Dla bezpośredniego modelu ILP "
        "oznacza to do 864 × 482 = 416 448 potencjalnych zmiennych położeń, "
        "dlatego ta instancja służy jako przypadek dla heurystyki, a nie jako "
        "praktyczny test dowodu optymalności."
    )
    document.add_paragraph(
        "Poprawność algorytmu dokładnego sprawdzono na małych instancjach, "
        "dla których optimum wyznaczono niezależnie przez pełny przegląd "
        "wszystkich sekwencji A/C/G/T zgodnych ze znanym początkiem. Testy "
        "uruchomiono w Pythonie 3.12.3 z biblioteką PuLP i programem CBC."
    )
    add_code_line(document, ".venv/bin/python -m unittest -v test_exact_b_m.py test_heur_b_m.py")

    document.add_heading("4.2 Wyniki walidacji algorytmu dokładnego", level=2)
    add_table(
        document,
        [
            "Przypadek",
            "n",
            "|S|",
            "Optimum z pełnego przeglądu",
            "Wynik ILP",
            "Czas ILP",
        ],
        [
            ("Pełne spektrum", 8, 5, 5, 5, "< 0,02 s"),
            ("Błędy negatywne", 8, 3, 3, 3, "< 0,02 s"),
            ("Symbole IUPAC", 7, 5, 5, 5, "< 0,02 s"),
            ("Konflikt nakładek", 7, 4, 4, 4, "< 0,02 s"),
        ],
        [4.1, 1, 1.1, 2.5, 1.8, 1.8],
    )
    document.add_paragraph(
        "We wszystkich czterech przypadkach wartość funkcji celu modelu ILP "
        "była równa optimum z pełnego przeglądu. Zestaw testów jednostkowych "
        "obejmuje ponadto parser XML; łącznie wykonano 6 testów."
    )
    document.add_paragraph(
        "Przypadek „Konflikt nakładek” jest testem regresyjnym. Poprzedni "
        "model grafowy oceniał zgodność sond tylko parami i dla tej instancji "
        "zwracał 2 pokryte sondy przy optimum równym 4. Bezpośredni model "
        "pozycji i nukleotydów osiąga wartość 4."
    )

    document.add_heading("4.3 Porównanie podejść", level=2)
    add_table(
        document,
        ["Kryterium", "Algorytm dokładny", "Heurystyka"],
        [
            (
                "Model",
                "ILP: nukleotydy i położenia sond",
                "Metoda zachłanna + ponowne uruchomienia + mutacje",
            ),
            (
                "Gwarancja",
                "Tak, wyłącznie przy statusie „Optimal”",
                "Brak gwarancji optimum",
            ),
            (
                "Walidacja",
                "Porównanie z pełnym przeglądem: 4/4 przypadki optymalne",
                "Test zachowania długości i znanego początku",
            ),
            (
                "Zastosowanie",
                "Małe instancje i punkt odniesienia",
                "Większe instancje",
            ),
        ],
        [3.2, 5.4, 5.4],
    )

    document.add_heading("4.4 Wnioski", level=2)
    add_manual_number(
        document,
        1,
        "Bezpośredni model ILP poprawnie uwzględnia jednoczesną zgodność "
        "nakładających się symboli IUPAC.",
    )
    add_manual_number(
        document,
        2,
        "Na małych instancjach model osiągnął optimum potwierdzone pełnym "
        "przeglądem, w tym dla danych z błędami negatywnymi.",
    )
    add_manual_number(
        document,
        3,
        "Status programu rozwiązującego musi być raportowany: tylko status "
        "„Optimal” oznacza dowód optymalności; najlepsze rozwiązanie znalezione "
        "przed upływem limitu czasu nie daje takiej gwarancji.",
    )
    add_manual_number(
        document,
        4,
        "Dla instancji n = 500 i 864 sond model dokładny jest zbyt duży do "
        "praktycznej walidacji na zwykłym komputerze, dlatego właściwym "
        "narzędziem pozostaje heurystyka.",
    )

    properties = document.core_properties
    properties.title = "Sekwencjonowanie przez hybrydyzację — sprawozdanie"
    properties.subject = "Binarny chip z błędami negatywnymi"
    properties.author = "Vasil Kusmartsev, Mateusz Kaźmierczak"

    document.save(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generuje sprawozdanie projektu w formacie DOCX."
    )
    parser.add_argument(
        "output",
        nargs="?",
        default=DEFAULT_OUTPUT,
        help=f"nazwa pliku wynikowego (domyślnie: {DEFAULT_OUTPUT})",
    )
    args = parser.parse_args()
    build_document(args.output)
